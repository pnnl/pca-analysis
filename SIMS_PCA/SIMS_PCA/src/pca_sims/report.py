"""This class is used for generating the report."""

import re
from typing import Dict, Any, Optional
import sys

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, Cm
from docx.shared import RGBColor

import pandas as pd
import numpy as np

class pca_sims_report(object):

    def __init__(self, f_report:str, ion_sign, description:Optional[Dict]=None):
        self.document = Document()
        self.f_report = f_report
        self.description = description

        # Set the page margins
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)

        # Title page
        self._create_title_page(ion_sign)
    
    
    def _create_title_page(self, ion_sign:str):
        document = self.document
        document.add_heading("\t\t\tPCA-SIMS Spectra Analysis Report", 0)

        # Add author and time
        description = self.description
        if description is not None:
            run=document.add_paragraph().add_run()
            run.add_break(); run.add_break();run.add_break()

            # p=document.add_paragraph("{}".format(description['experiment']))
            p=document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("{}".format(description['experiment']) + " (" + ion_sign + " ions)")
            run.bold = True
            run.font.size = Pt(18)
            run.add_break(); run.add_break();run.add_break()
            run.add_break(); run.add_break();run.add_break()

            # p=document.add_paragraph('ToF-SIMS testing date: {}'.format(description['date']))
            p=document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('ToF-SIMS testing date: {}'.format(description['date']))
            run.font.size = Pt(14)

            # p=document.add_paragraph('ToF-SIMS operator: {}'.format(description['operator']))
            p=document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('ToF-SIMS operator: {}'.format(description['operator']))
            run.font.size = Pt(14)
    

    def write_2dscore_plots(self, score2d_plots):
        document = self.document

        # Start a new page
        document.add_page_break()

        # Write the title
        document.add_heading('2D PCA scores plots')

        # Write the score plots
        for plots in score2d_plots:
            for plot in plots:
                document.add_picture(plot, width=Inches(6))  # (6" Score plots + 0.625" margin + 0.625" margin = 7.5" total)


    def write_plot_page(self, pcacomp:int, positive_ion:bool, score_plot:str, loading_plot:str,
                        positive_loading_table:pd.DataFrame, negative_loading_table:pd.DataFrame, 
                        ion_signals: Dict):
        sign_ion = 'Positive' if positive_ion else 'Negative'

        document = self.document
        # Start a new page
        document.add_page_break()

        # Write the title
        document.add_heading('{} ion spectra, PCA analysis results -- PC{}'.format(sign_ion, pcacomp), 1)

        # Write the score plots
        document.add_picture(score_plot, width=Inches(5))  # Score plot
        document.add_picture(loading_plot, width=Inches(5))  # Loading plot

        # Write the top 20 positive loading assignments
        document.add_paragraph("High score samples contain more:")
        p_pos = document.add_paragraph("", style="List Bullet")
        for index in positive_loading_table.index:
            unit_mass, assign = positive_loading_table.loc[index, "Unit Mass"], positive_loading_table.loc[index, "Initial Peak Assignment"]
            write_top_loadings_list(p_pos, unit_mass, assign)
            # Only add a comma separator if we are not on the last element
            if (index != positive_loading_table.index[-1]):
                p_pos.add_run(', ')
        # Remove the extra comma separator and space at the end of the run
        # p_pos.text = p_pos.text[:-2]
        positive_dominant_ions = [ion_type for ion_type in ion_signals if ion_signals[ion_type]['active'] and (ion_signals[ion_type]['type']=='+pca')] 
        document.add_paragraph(", ".join(positive_dominant_ions), style="List Bullet") # ion categories

        # Write the top 20 negative loading assignments
        document.add_paragraph("Low score samples contain more:")
        p_neg = document.add_paragraph("", style="List Bullet")
        for index in negative_loading_table.index:
            unit_mass, assign = negative_loading_table.loc[index, "Unit Mass"], negative_loading_table.loc[index, "Initial Peak Assignment"]
            write_top_loadings_list(p_neg, unit_mass, assign)
            # Only add a comma separator if we are not on the last element
            if (index != negative_loading_table.index[-1]):
                p_neg.add_run(', ')
        # Remove the extra comma separator and space at the end of the run
        # p_neg.text = p_neg.text[:-2]
        negative_dominant_ions = [ion_type for ion_type in ion_signals if ion_signals[ion_type]['active'] and (ion_signals[ion_type]['type']=='-pca')] 
        document.add_paragraph(", ".join(negative_dominant_ions), style="List Bullet") # ion categories


    # Writes a PCA loading table to the document. Columns are:
    # Loading No. | Unit Mass | Document Mass | Initial Peak Assignment | Initial Probabilities | Measured Mass | Peak Assignment (Qualified)
    #  | Updated Peak Assignment (from Document Mass) | Updated Document Mass
    def write_table_page(self, pcacomp:int, positive_ion:bool, 
                         p_loading_table:pd.DataFrame, n_loading_table:pd.DataFrame):
        document = self.document
        sign_ion = 'Positive' if positive_ion else 'Negative'

        # Start a new page
        document.add_page_break()

        # Write the title
        document.add_heading('{} ion spectra, top positive loadings -- PC{}'.format(sign_ion, pcacomp), 1)

        # Write positive loading values ...
        document_add_table(document, p_loading_table)
        
        # Start a new page
        document.add_page_break()

        # Write the title
        document.add_heading('{} ion spectra, top negative loadings -- PC{}'.format(sign_ion, pcacomp), 1)

        # Write negative loading values ...
        document_add_table(document, n_loading_table)


    def write_analysis_page(self, pcacomp:int, positive_ion:bool, 
                            positive_loading_table:pd.DataFrame, negative_loading_table:pd.DataFrame, 
                            ion_signals: Dict):
        document = self.document
        sign_ion = 'Positive' if positive_ion else 'Negative'

        # Start a new page
        document.add_page_break()

        # Write the title
        document.add_heading('{} ion spectra, molecular information from PC{} loadings plot'.format(sign_ion, pcacomp), 1)

        # Write the overview description
        p = document.add_paragraph("""The major positive PC{0} loadings are """.format(pcacomp), style="List Bullet")
        for index in positive_loading_table.index:
            unit_mass, assign = positive_loading_table.loc[index, "Unit Mass"], positive_loading_table.loc[index, "Initial Peak Assignment"]
            write_top_loadings_list(p, unit_mass, assign)
            p.add_run(', ')
        p.add_run("indicating they are more observed in high PC{0} score samples.".format(pcacomp))

        p = document.add_paragraph("""The major negative PC{0} loadings are """.format(pcacomp), style="List Bullet")
        for index in negative_loading_table.index:
            unit_mass, assign = negative_loading_table.loc[index, "Unit Mass"], negative_loading_table.loc[index, "Initial Peak Assignment"]
            write_top_loadings_list(p, unit_mass, assign)
            p.add_run(', ')
        p.add_run("indicating they are more observed in low PC{0} score samples.".format(pcacomp))

        # Write the dominant ion categories
        positive_dominant_ions = [ion_type for ion_type in ion_signals if ion_signals[ion_type]['active'] and (ion_signals[ion_type]['type']=='+pca')] 
        negative_dominant_ions = [ion_type for ion_type in ion_signals if ion_signals[ion_type]['active'] and (ion_signals[ion_type]['type']=='-pca')] 
        for ion_type in positive_dominant_ions:
            dominant_mass = ion_signals[ion_type]['top_ion_list']
            
            p = document.add_paragraph("{} signals, such as ".format(ion_type), style="List Bullet")
            for unit_mass in dominant_mass:
                assign = positive_loading_table.loc[unit_mass, "Initial Peak Assignment"]
                write_top_loadings_list(p, unit_mass, assign)
                p.add_run(', ')
            p.add_run("are mostly found in positive loadings, indicating that high PC{} score samples contain more {}.".format(pcacomp, ion_type))
        
        for ion_type in negative_dominant_ions:
            dominant_mass = ion_signals[ion_type]['top_ion_list']
            
            p = document.add_paragraph("{} signals, such as ".format(ion_type), style="List Bullet")
            for unit_mass in dominant_mass:
                assign = negative_loading_table.loc[unit_mass, "Initial Peak Assignment"]
                write_top_loadings_list(p, unit_mass, assign)
                p.add_run(', ')
            p.add_run("are mostly found in negative loadings, indicating that low PC{} score samples contain more {}.".format(pcacomp, ion_type))
    
    def save(self):
        # TODO: Generate table of contents

        # Add page numbers
        self.document.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document_add_page_number(self.document.sections[0].footer.paragraphs[0].add_run())

        # Save the report
        self.document.save(self.f_report)


# ------------------------------------------------------------------------------ Some useful helper methods ------------------------------------------------------------------------------

# Add a table to the end and create a reference variable along with an extra header row
def document_add_table(document:Document, df:pd.DataFrame):
    # Create number of rows + columns corresponding to the dataframe's size
    t = document.add_table(df.shape[0]+1, df.shape[1])
    t.style = 'Table Grid'

    # Add the headers at the top of each column
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]

    # Add the rest of the dataframe. Start at index 1 (see "i+1") to avoid conflicts with the header row, and handle 4 separate cases:
    # 1) Value is NaN - Add blank text to the cell
    # 2) Value is a list of non-floats - parse each string in the list (see document_add_assignment) and add it to a paragraph that is then added to the cell
    # 3) Value is a list of floats - convert each list entry to a string and add it to the cell
    # 4) Value is any non-list entity - convert it to a string and add it to the cell
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            cur_group = df.values[i,j]
            # We expect each group to be either a nested list of lists or a single-depth list, which we handle with the if case here; if list is empty, also default to the 1st case
            if not cur_group or not isinstance(cur_group, list) or not isinstance(cur_group[0], list):
                group_size = 1
            else:
                group_size = len(cur_group)

            # Iterate over multiple sublists of the current dataframe entry, but only if it has a nested list structure
            for k in range(group_size):
                if not cur_group or not isinstance(cur_group, list) or not isinstance(cur_group[0], list):
                    cur_entry = cur_group
                else:
                    cur_entry = cur_group[k]
                
                # TODO Adds an extra line at beginning of each cell; might want to remove this to shorten table?
                p = t.cell(i+1,j).add_paragraph()
                if str(cur_entry) == 'nan':                                                             # Do nothing for entries that are not a number
                    p.add_run('')
                elif isinstance(cur_entry, list) and (not cur_group or not is_float(cur_entry[0])):     # Lists of species classifications
                    document_add_assignment(p, cur_entry, in_table=True)
                elif isinstance(cur_entry, list) and (not cur_group or is_float(cur_entry[0])):         # Probabilities and lists of document masses
                    p.add_run('\n'.join(map(str, cur_entry)))
                else:                                                                                   # Anything else
                    p.add_run(str(cur_entry))
                
                # TODO Why are we checking t.cell(i+1,j).text.strip()? Is j wrong? Could new lines be fooling this into running every row anyway?
                # Change text color to yellow if deviation is 100-200ppm, red if deviation is > 200ppm, and green if < 100ppm. This should give a user
                # a simple indicator to tell if they need to check the vailidity of the data or the peak assignments. Also, check the column header and 
                # only apply highlighting we are in the column 'Peak Assignment (Qualified)' (and if there is actually a value in the current cell of 
                # that column).
                header_text = t.cell(0,j).text
                if header_text == 'Peak Assignment (Qualified)' and t.cell(i+1,j).text.strip():
                    # Calculate the deviations between the measured masses and document masses and express them as floats; use these to highlight large errors.
                    # TODO Note that cur_doc_masses is the old entry and cur_updated_doc_mass is the newest one. Currently, using newer entry to verify past measurements,
                    # but data could be out of order and may have to use old entry according to others; change later if needed.
                    # TODO Currently cutting off any elements beyond first n (i.e., if one array is longer than the other); may need to fix in future.
                    try:
                        # Save the current measured mass (1 cell to the left) and list of document masses (4 cells to the left). Make sure to transform the doc
                        # masses into a NumPy array of floats, from which we can get the list of deviations between the measured mass and doc masses by broadcasting.
                        cur_measured_masses = t.cell(i+1,j-1).text
                        cur_measured_masses = cur_measured_masses.strip()
                        cur_measured_masses = re.split('[,\n]+', cur_measured_masses)
                        cur_measured_masses = np.array(cur_measured_masses)
                        cur_measured_masses = cur_measured_masses.astype(float)

                        cur_doc_masses = t.cell(i+1,j-4).text
                        cur_doc_masses = cur_doc_masses.strip()
                        cur_doc_masses = re.split('[,\n]+', cur_doc_masses)
                        cur_doc_masses = np.array(cur_doc_masses)
                        cur_doc_masses = cur_doc_masses.astype(float)
                    except:
                        print("Error! Encountered row missing a Document Mass entry. Please fix the report before trying again.")
                        sys.exit()

                    # The measured masses array may be shorter than the document masses array; if so, just calculate the deviations for a number of 
                    # elements equal to that found in the measured masses array. However, if there is only one measured mass entry, use that to calculate 
                    # deviations and apply highlighting to all peak assignments. If the measured masses array is somehow larger than the document masses
                    # array, throw an exception and tell the user to fix the report before running again.
                    try:
                        n_mm = len(cur_measured_masses)
                        n_dm = len(cur_doc_masses)

                        if n_mm == 1:
                            # Subtract single MM entry from all DM entries
                            cur_deviations_array = abs(cur_doc_masses[:n_dm] - cur_measured_masses)
                            cur_fractional_deviations_array = cur_deviations_array / cur_measured_masses
                        else:
                            # Otherwise, calculate deviations normally
                            cur_deviations_array = abs(cur_doc_masses[:n_mm] - cur_measured_masses[:n_mm])
                            cur_fractional_deviations_array = cur_deviations_array / cur_measured_masses
                    except:
                        print("Error! There are more measured masses than document masses in the row containing " + 
                               "the following measured mass values: ", cur_measured_masses, "\n")
                        print("Please fix the number of measured masses in this row, update the database, and run again.")
                        sys.exit()

                    
                    # Iterate over each line in the qualified peak assignments column, and using the corresponding document masses along with their
                    # deviations from the measured masses, highlight each in the correct font color. If there is only one measured mass, we will highlight
                    # each line according to the deviations from that mass to make sure each possible peak assignment is qualified.
                    paragraph_lines = p.text.split("\n")
                    runs_ind = 0
                    line_ind = 0
                    for l in range(len(cur_fractional_deviations_array)):   # Use the number of deviations calculated above to determine how many lines we have to highlight
                        cur_fdev = cur_fractional_deviations_array[l]

                        # Set the current font color based on the current fractional deviation
                        if cur_fdev < 0.0001:
                            cur_color = RGBColor(0, 220, 50)         # Green
                        elif cur_fdev > 0.0002:
                            cur_color = RGBColor(255, 0, 0)         # Red
                        else:
                            cur_color = RGBColor(220, 165, 0)       # Yellow
                        
                        
                        # Iterate over each run in the current line of the paragraph and change the font color of each to the current color.
                        # For paragraphs where there is no new line (i.e., cells with only one entry), exit when we hit the last index of p.runs.
                        while line_ind <= l and runs_ind < len(p.runs):
                            cur_run = p.runs[runs_ind]

                            # Increment line index counter every time we hit a new line. Lines may have multiple runs, so this makes sure we get the
                            # correct number of lines by only counting up on the new line character that signifies the end of each line.
                            if cur_run.text == '\n':
                                line_ind += 1

                            # Only change fonts if we have arrived at the current line
                            if line_ind == l:
                                cur_run.font.color.rgb = cur_color

                            runs_ind += 1


# TODO Superscripting is done wrong on - sign after a number and for ? marks (for example: see SNO2-?)
# Parse strings representing the species assignments that we intend to add to a cell.
#   Parameters:
#   p - (docx paragraph) The paragraph to which we add text
#   assignment - The list of items to add to the paragraph
#   in_table - True means we use new line separators to add text to a table, False means we use comma separators to add text elsewhere
def document_add_assignment(p, assignment:list, in_table:bool=False) -> None:
    n_assign = len(assignment)
    # For each chemical species assignment (expect assign to be a string), pull off sign(s) on the end, then split the remaining string on any digits in the
    # chemical formula (e.g., 'C3H6' -> ['C', '3', 'H', '6']); afterward, recombine this string and its charge with proper subscripting and superscripting 
    # for chemical formulas
    for j,assign in enumerate(assignment):
        # Tell the user if a charge, which should be indicated by some number of + or - signs at the end of each assign, is missing on one of the entries
        try:
            assign_, signs = re.split('([+-]+)', assign)[0:2]
            assign_ = re.split('(\d+)', assign_)
            assign_ = [s for s in assign_ if s != '']
        except:
            print('***Error! Ion missing charge! Please ensure each species you entered has at least one + or - sign at the end.***\n')
            print("Species causing error: ", assign)
            sys.exit()

        # Add number and letter
        for i,s in enumerate(assign_):
            if i==0 and is_float(s):
                text = p.add_run(s)
                text.font.superscript = True
            elif i!=0 and is_float(s):
                text = p.add_run(s)
                text.font.subscript = True
            else:
                text = p.add_run(s)
        # Add any sign(s) at the end
        text = p.add_run(signs)
        text.font.superscript = True
        # If not the last element, add a separator appropriate to the context (table or not in table)
        if j != n_assign-1 and in_table:
            p.add_run('\n')
        elif j != n_assign-1:
            p.add_run(', ')
        

def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def document_add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def is_float(element: Any) -> bool:
    try:
        float(element)
        return True
    except ValueError:
        return False
    

# Encapsulates boilerplate for writing the loadings lists in bullet point form
def write_top_loadings_list(p, unit_mass: float, assign: str):
    # Add all m/z ratios and their corresponding assignments (if we can get them)
    if str(assign) == 'nan':
        p.add_run("m/z {} (-)".format(unit_mass))
    else:
        p.add_run("m/z {} (".format(unit_mass))
        document_add_assignment(p, assign)
        p.add_run(")")